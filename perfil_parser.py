"""
Extrae texto de un perfil de cargo (.docx o .pdf) y hace una primera pasada de
extracción de requisitos (formación, palabras clave de RRLL/sindicatos,
palabras clave de industria y rango salarial).

Esta extracción es heurística (búsqueda de palabras conocidas + regex de
montos). No reemplaza la revisión humana: la interfaz permite editar los
resultados antes de correr el filtro.
"""

import io
import re

import docx
import pdfplumber


def extract_text_from_docx(file_stream):
    document = docx.Document(file_stream)
    lines = []
    for p in document.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text_from_pdf(file_stream):
    lines = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.append(text)
    return "\n".join(lines)


def extract_perfil_text(filename, file_bytes):
    stream = io.BytesIO(file_bytes)
    lower = filename.lower()
    if lower.endswith(".docx"):
        return extract_text_from_docx(stream)
    elif lower.endswith(".pdf"):
        return extract_text_from_pdf(stream)
    else:
        raise ValueError(
            "Formato no soportado (%s). Usa un archivo .docx o .pdf" % filename
        )


# Carreras / formaciones típicas que suelen aparecer como excluyentes en
# perfiles de RRHH, agrupadas por "familia" (forma carrera / forma profesión,
# masculino / femenino). Si el perfil de cargo menciona cualquier variante de
# una familia, se acepta como cumplida cualquier otra variante de esa misma
# familia en el CV del candidato (ej. el perfil dice "Psicología" y el CV del
# candidato dice "Psicólogo" -> igual cuenta como match).
CAREER_FAMILIES = [
    ["ingeniería comercial", "ingeniero comercial", "ingeniera comercial"],
    [
        "ingeniería civil industrial",
        "ingeniero civil industrial",
        "ingeniera civil industrial",
    ],
    ["psicología", "psicólogo", "psicóloga"],
    [
        "contador auditor",
        "contadora auditora",
        "contador general",
        "contadora general",
    ],
    [
        "ingeniería en recursos humanos",
        "ingeniero en recursos humanos",
        "ingeniera en recursos humanos",
    ],
    [
        "administración de empresas",
        "administrador de empresas",
        "administradora de empresas",
        "ingeniería en administración de empresas",
    ],
    [
        "administración de personal",
        "administrador de personal",
        "administradora de personal",
    ],
]

# Lista plana (todas las variantes), usada solo para buscar en el texto del
# perfil de cargo.
KNOWN_CAREERS = [variant for family in CAREER_FAMILIES for variant in family]

DEFAULT_RRLL_KEYWORDS = [
    "sindicato",
    "sindical",
    "negociación colectiva",
    "relaciones laborales",
    "convenio colectivo",
    "dirigentes sindicales",
]

DEFAULT_INDUSTRIA_KEYWORDS = [
    "industrial",
    "planta productiva",
    "manufactura",
    "producción",
    "alimentos",
    "consumo masivo",
    "forestal",
    "salmonera",
    "minería",
    "agroindustria",
    "logística",
]

SALARY_RE = re.compile(r"\$?\s?([\d][\d\.]{5,10})")

# Líneas que probablemente hablan de la formación/estudios requeridos (para
# no buscar nombres de carrera en cualquier parte del documento, solo donde
# tiene sentido).
FORMACION_TRIGGER_WORDS = [
    "formación",
    "formacion",
    "estudios",
    "profesión",
    "profesion",
    "carrera",
    "requisitos académicos",
    "requisitos academicos",
    "título profesional",
    "titulo profesional",
    "nivel educacional",
    "educación",
    "educacion",
]

# Patrones genéricos de nombres de carrera, para detectar formaciones que NO
# están en nuestra lista fija CAREER_FAMILIES (ej. "Ingeniería Industrial",
# "Ingeniería en Alimentos", "Técnico en Prevención de Riesgos"). Esto evita
# tener que hardcodear cada carrera posible: cualquier perfil de cargo nuevo
# queda cubierto mientras use estos patrones comunes de redacción.
GENERIC_CAREER_RE = re.compile(
    r"(?:"
    r"ingenier[íi]a?\s+(?:civil\s+)?(?:en\s+|de\s+)?[a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){0,2}"
    r"|licenciatura\s+en\s+[a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){0,2}"
    r"|licenciad[oa]\s+en\s+[a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){0,2}"
    r"|técnic[oa]\s+(?:de\s+nivel\s+superior\s+)?en\s+[a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){0,2}"
    r"|contador(?:a)?(?:\s+(?:auditor(?:a)?|general))?"
    r"|administrad(?:or|ora)\s+(?:de|en)\s+[a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){0,2}"
    r"|psicólog[oa]"
    r")",
    re.IGNORECASE,
)


ENUMERATION_SPLIT_RE = re.compile(r",|;|\so\s|\su\s|\sy\s", re.IGNORECASE)


def _extract_generic_formacion(full_text):
    found = []
    for line in full_text.split("\n"):
        line_lower = line.lower()
        if not any(tw in line_lower for tw in FORMACION_TRIGGER_WORDS):
            continue
        # Se separa por comas/"o"/"y" antes de aplicar el regex, para que al
        # buscar una carrera no se "coma" de paso las primeras palabras de
        # la siguiente carrera enumerada en la misma línea (ej. "Ingeniería
        # Industrial, Ingeniería en Alimentos o carreras afines").
        for fragment in ENUMERATION_SPLIT_RE.split(line):
            for match in GENERIC_CAREER_RE.findall(fragment):
                cleaned = match.strip().lower()
                if cleaned and cleaned not in found:
                    found.append(cleaned)
    return found


def _extract_salary_range(full_text):
    matches = SALARY_RE.findall(full_text)
    values = []
    for m in matches:
        digits = re.sub(r"\.", "", m)
        if digits.isdigit():
            val = int(digits)
            if 500_000 <= val <= 20_000_000:
                values.append(val)
    if not values:
        return None, None
    return min(values), max(values)


def parse_requirements(full_text):
    text_lower = full_text.lower()

    formacion_excluyente = []
    for family in CAREER_FAMILIES:
        if any(variant in text_lower for variant in family):
            for variant in family:
                if variant not in formacion_excluyente:
                    formacion_excluyente.append(variant)

    # Respaldo genérico: si el cargo pide una carrera que no está en nuestra
    # lista conocida (CAREER_FAMILIES), igual queda detectada buscando
    # patrones de nombres de carrera en las líneas que mencionan
    # formación/estudios/requisitos académicos.
    for extra in _extract_generic_formacion(full_text):
        if extra not in formacion_excluyente:
            formacion_excluyente.append(extra)

    salario_min, salario_max = _extract_salary_range(full_text)

    # Igual que con la formación: las palabras clave de RRLL/sindicatos e
    # industria solo se autocompletan si el PROPIO perfil de cargo las
    # menciona. La mayoría de los cargos no tienen relación con sindicatos
    # ni con una industria específica, y llenar esos campos "por defecto"
    # sin que el perfil lo pida le sumaba puntaje a candidatos por
    # coincidencias irrelevantes. Si el cargo sí lo requiere, quedan
    # detectadas automáticamente; si no, el campo queda vacío (y se puede
    # completar a mano en la interfaz si aplica).
    rrll_relevante = any(kw in text_lower for kw in DEFAULT_RRLL_KEYWORDS)
    industria_relevante = any(kw in text_lower for kw in DEFAULT_INDUSTRIA_KEYWORDS)

    return {
        "formacion_excluyente": formacion_excluyente,
        "rrll_keywords": list(DEFAULT_RRLL_KEYWORDS) if rrll_relevante else [],
        "industria_keywords": list(DEFAULT_INDUSTRIA_KEYWORDS)
        if industria_relevante
        else [],
        "salario_min": salario_min,
        "salario_max": salario_max,
        "raw_text_preview": full_text[:4000],
    }
