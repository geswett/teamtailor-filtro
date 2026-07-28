"""
Extrae texto de un perfil de cargo (.docx o .pdf) y hace una primera pasada de
extracción de requisitos (formación, palabras clave de RRLL/sindicatos,
palabras clave de industria y rango salarial).

Esta extracción es heurística (búsqueda de palabras conocidas + regex de
montos). No reemplaza la revisión humana: la interfaz permite editar los
resultados antes de correr el filtro.
"""

import io
import re

import docx
import pdfplumber


def extract_text_from_docx(file_stream):
    document = docx.Document(file_stream)
    lines = []
    for p in document.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text_from_pdf(file_stream):
    lines = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.append(text)
    return "\n".join(lines)


def extract_perfil_text(filename, file_bytes):
    stream = io.BytesIO(file_bytes)
    lower = filename.lower()
    if lower.endswith(".docx"):
        return extract_text_from_docx(stream)
    elif lower.endswith(".pdf"):
        return extract_text_from_pdf(stream)
    else:
        raise ValueError(
            "Formato no soportado (%s). Usa un archivo .docx o .pdf" % filename
        )


# Carreras / formaciones típicas que suelen aparecer como excluyentes en
# perfiles de RRHH, agrupadas por "familia" (forma carrera / forma profesión,
# masculino / femenino). Si el perfil de cargo menciona cualquier variante de
# una familia, se acepta como cumplida cualquier otra variante de esa misma
# familia en el CV del candidato (ej. el perfil dice "Psicología" y el CV del
# candidato dice "Psicólogo" -> igual cuenta como match).
CAREER_FAMILIES = [
    ["ingeniería comercial", "ingeniero comercial", "ingeniera comercial"],
    [
        "ingeniería civil industrial",
        "ingeniero civil industrial",
        "ingeniera civil industrial",
    ],
    ["psicología", "psicólogo", "psicóloga"],
    [
        "contador auditor",
        "contadora auditora",
        "contador general",
        "contadora general",
    ],
    [
        "ingeniería en recursos humanos",
        "ingeniero en recursos humanos",
        "ingeniera en recursos humanos",
    ],
    [
        "administración de empresas",
        "administrador de empresas",
        "administradora de empresas",
        "ingeniería en administración de empresas",
    ],
    [
        "administración de personal",
        "administrador de personal",
        "administradora de personal",
    ],
]

# Lista plana (todas las variantes), usada solo para buscar en el texto del
# perfil de cargo.
KNOWN_CAREERS = [variant for family in CAREER_FAMILIES for variant in family]

DEFAULT_RRLL_KEYWORDS = [
    "sindicato",
    "sindical",
    "negociación colectiva",
    "relaciones laborales",
    "convenio colectivo",
    "dirigentes sindicales",
]

DEFAULT_INDUSTRIA_KEYWORDS = [
    "industrial",
    "planta productiva",
    "manufactura",
    "producción",
    "alimentos",
    "consumo masivo",
    "forestal",
    "salmonera",
    "minería",
    "agroindustria",
    "logística",
]

SALARY_RE = re.compile(r"\$?\s?([\d][\d\.]{5,10})")


def _extract_salary_range(full_text):
    matches = SALARY_RE.findall(full_text)
    values = []
    for m in matches:
        digits = re.sub(r"\.", "", m)
        if digits.isdigit():
            val = int(digits)
            if 500_000 <= val <= 20_000_000:
                values.append(val)
    if not values:
        return None, None
    return min(values), max(values)


def parse_requirements(full_text):
    text_lower = full_text.lower()

    formacion_excluyente = []
    for family in CAREER_FAMILIES:
        if any(variant in text_lower for variant in family):
            for variant in family:
                if variant not in formacion_excluyente:
                    formacion_excluyente.append(variant)

    salario_min, salario_max = _extract_salary_range(full_text)

    return {
        "formacion_excluyente": formacion_excluyente,
        "rrll_keywords": list(DEFAULT_RRLL_KEYWORDS),
        "industria_keywords": list(DEFAULT_INDUSTRIA_KEYWORDS),
        "salario_min": salario_min,
        "salario_max": salario_max,
        "raw_text_preview": full_text[:4000],
    }
